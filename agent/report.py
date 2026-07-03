"""Export an analysis as a regulated-style Statistical Analysis Report (.docx).

Renders an AgentResult into a Word document with the sections a reviewer expects: objective,
methods, the full data-engineering audit trail, results, assumption diagnostics, interpretation,
and an explicit limitations/validation statement. Honest by construction — it documents that this
is an automated exploratory analysis, not a validated pre-specified regulatory submission.
"""
from __future__ import annotations

import datetime as _dt
import io


def _kv(doc, key, val):
    p = doc.add_paragraph()
    p.add_run(f"{key}: ").bold = True
    p.add_run(str(val))


def _split_notes(issues):
    """Separate data-preparation steps from model/assumption diagnostics for the report."""
    prep, diag = [], []
    for i in issues or []:
        low = i.lower()
        if low.startswith(("dropped", "imputed", "removed")):
            prep.append(i)
        else:
            diag.append(i)
    return prep, diag


def _mono(p):
    from docx.shared import Pt
    for r in p.runs:
        r.font.name = "Consolas"
        r.font.size = Pt(9)


def _chart_png(chart):
    """Render an Altair chart to PNG bytes (vl-convert); None if it can't be rendered."""
    if chart is None:
        return None
    try:
        buf = io.BytesIO()
        chart.save(buf, format="png", ppi=120)
        return buf.getvalue()
    except Exception:
        return None


def _figures_for(result, m):
    """The Altair figure(s) that match the on-screen result, so the report shows the same charts."""
    from agent import charts as ch
    mt = m.get("model_type")
    figs = []
    try:
        if mt == "timeseries":
            figs = [ch.forecast_chart(m.get("series"))]
        elif mt == "forest":
            figs = [ch.importance_chart(m)]
        elif mt == "experiment":
            figs = [ch.experiment_chart(m)]
        elif mt == "noninferiority":
            figs = [ch.ni_plot(m)]
        elif mt == "sample_size":
            figs = [ch.power_curve_chart(m)]
        elif mt == "survival":
            figs = [ch.survival_plot(m.get("km")), ch.forest_plot(m)]
        elif m.get("terms"):                         # logistic / ols / cox / causal / association
            figs = [ch.forest_plot(m)]
        elif not m and getattr(result, "dataframe", None) is not None:
            figs = [ch.build_chart(result.dataframe)]
    except Exception:
        figs = []
    return [f for f in figs if f is not None]


def build_docx(result, *, when: _dt.datetime | None = None) -> bytes:
    """Build the .docx report for an AgentResult and return its bytes."""
    from docx import Document

    ts = (when or _dt.datetime.now()).strftime("%Y-%m-%d %H:%M")
    doc = Document()
    doc.add_heading("Statistical Analysis Report", 0)
    sub = doc.add_paragraph()
    sub.add_run(f"Clinical Insight Agent · OpenAI · generated {ts}").italic = True
    doc.add_paragraph(
        "Synthetic / illustrative data — not clinical fact. Automated EXPLORATORY analysis with "
        "transparent data engineering and assumption diagnostics; NOT a validated, pre-specified, "
        "double-programmed regulatory analysis.")

    m = result.model or {}

    doc.add_heading("1. Objective", 1)
    doc.add_paragraph(result.question or "")
    if getattr(result, "hypothesis", None):
        doc.add_heading("2. Hypothesis", 1)
        doc.add_paragraph(result.hypothesis)

    doc.add_heading("3. Data & Methods", 1)
    if m.get("model_type"):
        _kv(doc, "Method", m["model_type"].upper())
    if m.get("fit_stat"):
        _kv(doc, "Fit / design", m["fit_stat"])
    if m.get("n") is not None:
        _kv(doc, "Analysis n", f"{m.get('n', 0):,}")
    if getattr(result, "citations", None):
        _kv(doc, "Source tables", ", ".join(result.citations))
    if getattr(result, "sql", None):
        doc.add_paragraph("Analytic query:").runs[0].bold = True
        _mono(doc.add_paragraph(result.sql))

    prep, diag = _split_notes(m.get("issues"))
    if prep:
        doc.add_heading("3.1 Data preparation (audit trail)", 2)
        for s in prep:
            doc.add_paragraph(s, style="List Bullet")

    doc.add_heading("4. Results", 1)
    if m.get("verdict"):
        _kv(doc, "Conclusion", m["verdict"].get("call", ""))
        if m["verdict"].get("reason"):
            doc.add_paragraph(m["verdict"]["reason"])

    if m.get("model_type") == "sample_size":
        for a in m.get("arms", []):
            _kv(doc, a["arm"], f"{a['n']:,} subjects")
    elif m.get("model_type") == "timeseries":       # forecast output, not the raw input series
        series = m.get("series", [])
        hist = [p for p in series if p.get("kind") == "history"]
        fc = [p for p in series if p.get("kind") == "forecast"]
        if hist:
            _kv(doc, "Last observed", f"{hist[-1]['time'][:10]} = {hist[-1]['value']:.1f}")
        t = doc.add_table(rows=1, cols=3)
        t.style = "Table Grid"
        for j, h in enumerate(["Period", "Forecast", "95% band"]):
            t.rows[0].cells[j].text = h
        for p in fc:
            c = t.add_row().cells
            c[0].text, c[1].text = p["time"][:10], f"{p['value']:.1f}"
            c[2].text = f"[{p['lower']:.1f}, {p['upper']:.1f}]"
    elif m.get("arms"):
        binm = all(0 <= a["value"] <= 1 for a in m["arms"])
        t = doc.add_table(rows=1, cols=4)
        t.style = "Table Grid"
        for j, h in enumerate(["Arm", "Estimate", "95% CI", "n"]):
            t.rows[0].cells[j].text = h
        for a in m["arms"]:
            c = t.add_row().cells
            val = f"{a['value'] * 100:.1f}%" if binm else f"{a['value']:.3f}"
            ci = ("—" if a["ci_low"] != a["ci_low"] else
                  (f"[{a['ci_low'] * 100:.1f}%, {a['ci_high'] * 100:.1f}%]" if binm
                   else f"[{a['ci_low']:.3f}, {a['ci_high']:.3f}]"))
            c[0].text, c[1].text, c[2].text, c[3].text = a["arm"], val, ci, f"{a['n']:,}"
    elif m.get("terms"):
        t = doc.add_table(rows=1, cols=4)
        t.style = "Table Grid"
        for j, h in enumerate(["Term", m.get("effect_label", "estimate"), "95% CI", "p-value"]):
            t.rows[0].cells[j].text = h
        for term in m["terms"]:
            c = t.add_row().cells
            lo = term["ci_low"]
            ci = "—" if lo != lo else f"[{lo:.3f}, {term['ci_high']:.3f}]"
            pv = term["p"]
            pcol = "—" if pv != pv else f"{pv:.4f}"
            c[0].text, c[1].text, c[2].text, c[3].text = term["name"], f"{term['estimate']:.3f}", ci, pcol
    elif getattr(result, "dataframe", None) is not None:
        df = result.dataframe.head(30)
        t = doc.add_table(rows=1, cols=len(df.columns))
        t.style = "Table Grid"
        for j, col in enumerate(df.columns):
            t.rows[0].cells[j].text = str(col)
        for _, row in df.iterrows():
            c = t.add_row().cells
            for j, col in enumerate(df.columns):
                c[j].text = str(row[col])

    from docx.shared import Inches
    for i, fig in enumerate(_figures_for(result, m)):
        png = _chart_png(fig)
        if png:
            doc.add_paragraph(f"Figure {i + 1}.").runs[0].italic = True
            doc.add_picture(io.BytesIO(png), width=Inches(6.0))

    if diag:
        doc.add_heading("5. Assumptions & diagnostics", 1)
        for s in diag:
            doc.add_paragraph(s, style="List Bullet")
    if getattr(result, "findings", None):
        doc.add_heading("5.1 Statistical guardrail", 1)
        for f in result.findings:
            doc.add_paragraph(f"[{f.severity.upper()}] {f.kind} — {f.message}", style="List Bullet")

    if getattr(result, "interpretation", None):
        doc.add_heading("6. Interpretation", 1)
        for line in result.interpretation.split("\n"):
            if line.strip():
                doc.add_paragraph(line.replace("**", "").lstrip("# ").strip())

    doc.add_heading("7. Limitations & validation statement", 1)
    doc.add_paragraph(
        "Data are synthetic and illustrative; magnitudes are not empirical. Variable selection here is "
        "data-driven (collinearity/quasi-constant removal), which is appropriate for exploration but NOT "
        "for a confirmatory analysis, where the analysis set and methods must be pre-specified in a SAP. "
        "This report is machine-generated and has not been independently double-programmed or validated. "
        "A qualified biostatistician must review it before any decision or regulatory use.")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
